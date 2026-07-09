# -*- coding: utf-8 -*-
"""
AIZS 项目介绍书 Markdown 转 Word 转换工具
"""
import os
import re
import sys
from pathlib import Path

# 确保 docx 库可用
try:
    from docx import Document
    from docx.shared import Inches, Pt
except ImportError:
    print("未安装 python-docx 库，跳过 Word 生成。")
    sys.exit(0)

def convert_md_to_docx(md_path, docx_path):
    print(f"正在将 {md_path} 转换为 {docx_path}...")
    doc = Document()
    
    # 设定文档格式与页边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # 基础样式设定
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft YaHei'
    font.size = Pt(11)

    if not os.path.exists(md_path):
        print(f"找不到源文件: {md_path}")
        return

    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    in_code_block = False
    in_table = False
    table_rows = []

    for line in lines:
        line_stripped = line.strip()
        
        # 处理代码块
        if line_stripped.startswith('```'):
            in_code_block = not in_code_block
            continue

        if in_code_block:
            # 代码块段落样式设定
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            run = p.add_run(line)
            run.font.name = 'Consolas'
            run.font.size = Pt(9.5)
            continue

        # 处理表格
        if line_stripped.startswith('|') and line_stripped.endswith('|'):
            # 如果是表格分隔线如 |:---|:---|
            if '---' in line_stripped:
                continue
            
            in_table = True
            cells = [c.strip() for c in line_stripped.split('|')[1:-1]]
            table_rows.append(cells)
            continue
        else:
            if in_table and table_rows:
                # 渲染收集到的表格数据
                cols_count = max(len(r) for r in table_rows)
                table = doc.add_table(rows=len(table_rows), cols=cols_count)
                table.style = 'Table Grid'
                for r_idx, row_cells in enumerate(table_rows):
                    for c_idx, val in enumerate(row_cells):
                        if c_idx < len(table.rows[r_idx].cells):
                            # 清除 Markdown 强调符号
                            val_clean = val.replace('**', '').replace('`', '').replace('<br>', '\n')
                            table.rows[r_idx].cells[c_idx].text = val_clean
                doc.add_paragraph() # 空白段落分隔
                table_rows = []
                in_table = False

        if not line_stripped:
            continue

        # 处理标题
        if line_stripped.startswith('# '):
            doc.add_heading(line_stripped[2:], level=0)
        elif line_stripped.startswith('## '):
            doc.add_heading(line_stripped[3:], level=1)
        elif line_stripped.startswith('### '):
            doc.add_heading(line_stripped[4:], level=2)
        elif line_stripped.startswith('#### '):
            doc.add_heading(line_stripped[5:], level=3)
        # 处理无序列表
        elif line_stripped.startswith('* ') or line_stripped.startswith('- '):
            clean_text = line_stripped[2:].replace('**', '').replace('`', '')
            # 简单剥离 markdown 链接格式 [text](url) -> text
            clean_text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', clean_text)
            doc.add_paragraph(clean_text, style='List Bullet')
        # 处理常规段落
        else:
            clean_text = line_stripped.replace('**', '').replace('`', '')
            clean_text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', clean_text)
            doc.add_paragraph(clean_text)

    doc.save(docx_path)
    print(f"成功生成 Word 文档: {docx_path}")

if __name__ == '__main__':
    md_file = Path(__file__).resolve().parent.parent / 'docs' / 'AIZS_项目介绍书.md'
    docx_file = Path(__file__).resolve().parent.parent / 'docs' / 'AIZS_项目介绍书.docx'
    convert_md_to_docx(str(md_file), str(docx_file))
