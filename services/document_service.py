import os
import pdfplumber
import pypdf
import docx
from typing import List, Dict, Any, Optional
from utils.text_splitter import ChineseTextSplitter
from utils.logger import logger

class DocumentService:
    def __init__(self, chunk_size: int = 500, chunk_overlap: int = 100):
        """
        初始化文档解析服务。
        """
        self.splitter = ChineseTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)

    def parse_document(self, file_path: str, file_type: str) -> List[Dict[str, Any]]:
        """
        根据文件类型解析本地文件，并切分为结构化数据块。
        返回切片列表，每个切片格式为：
        {
          "text": str,
          "page_number": int or None,  # PDF 为 1-based 页码，DOCX/TXT 为 None
          "chunk_index": int           # 全局切片序号 (0-based)
        }
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"未找到待解析文件: {file_path}")

        file_type = file_type.lower().strip('.')
        logger.info(f"开始解析文件: {file_path}, 类型: {file_type}")

        if file_type == "pdf":
            return self._parse_pdf(file_path)
        elif file_type == "docx":
            return self._parse_docx(file_path)
        elif file_type in ["txt", "text"]:
            return self._parse_txt(file_path)
        else:
            raise ValueError(f"不支持的文件解析类型: {file_type}")

    def _parse_pdf(self, file_path: str) -> List[Dict[str, Any]]:
        """
        解析 PDF 文件，按页码进行段落切片以保留物理页码信息。
        优先使用 pdfplumber（兼容性更强），失败时回退到 pypdf。
        """


        # 优先尝试 pdfplumber（对复杂 PDF、嵌入式字体等兼容性更好）
        try:
            with pdfplumber.open(file_path) as pdf:
                total_pages = len(pdf.pages)
                logger.info(f"PDF 文件共有 {total_pages} 页（pdfplumber）")

                chunks, empty_page_count = self._process_pdf_pages(pdf.pages)

            logger.info(f"PDF 解析完成（pdfplumber），共切分为 {len(chunks)} 个切片")

            # 如果 pdfplumber 也没提取到文本，回退到 pypdf 再试一次
            if not chunks:
                logger.info(f"pdfplumber 未提取到文本，尝试 pypdf 回退解析")
                chunks, empty_page_count = self._parse_pdf_with_pypdf(file_path)

        except Exception as e:
            logger.warning(f"pdfplumber 解析失败: {e}，回退到 pypdf")
            chunks, empty_page_count = self._parse_pdf_with_pypdf(file_path)

        if not chunks:
            logger.warning(f"PDF 文件解析后无有效文本切片: {file_path}")

        return chunks

    def _parse_pdf_with_pypdf(self, file_path: str) -> tuple:
        """
        使用 pypdf 回退解析 PDF 文件。
        """

        try:
            reader = pypdf.PdfReader(file_path)
            total_pages = len(reader.pages)
            logger.info(f"PDF 文件共有 {total_pages} 页（pypdf 回退）")

            chunks, empty_page_count = self._process_pdf_pages(reader.pages)

            logger.info(f"PDF 解析完成（pypdf），共切分为 {len(chunks)} 个切片")
        except Exception as e:
            logger.error(f"pypdf 解析也失败: {e}")
            raise e
        return chunks, empty_page_count


    def _process_pdf_pages(self, pages) -> tuple:
        """
        提取并处理 PDF 页面内容的公共方法。
        返回 (chunks, empty_page_count)
        """
        chunks = []
        chunk_index = 0
        empty_page_count = 0

        for i, page in enumerate(pages):
            page_number = i + 1
            page_text = page.extract_text()
            if not page_text or not page_text.strip():
                empty_page_count += 1
                logger.debug(f"PDF 第 {page_number} 页无文本内容，跳过")
                continue

            page_chunks = self.splitter.split_text(page_text)
            for pc in page_chunks:
                chunks.append({
                    "text": pc,
                    "page_number": page_number,
                    "chunk_index": chunk_index
                })
                chunk_index += 1

        return chunks, empty_page_count

    def _parse_docx(self, file_path: str) -> List[Dict[str, Any]]:
        """
        解析 Word (docx) 文件。因为 Word 物理页码依赖排版渲染，此处没有可靠页码，置为 None 并记录片段编号。
        """
        chunks = []
        try:
            doc = docx.Document(file_path)
            paragraphs_text = []
            
            # 提取段落内容
            for p in doc.paragraphs:
                text = p.text.strip()
                if text:
                    paragraphs_text.append(text)
            
            # 提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if row_text:
                        paragraphs_text.append(row_text)

            full_text = "\n".join(paragraphs_text)
            if not full_text.strip():
                logger.warning(f"DOCX 文件无有效文本内容: {file_path}")
                return []

            # 切分全文
            split_chunks = self.splitter.split_text(full_text)
            for i, c in enumerate(split_chunks):
                chunks.append({
                    "text": c,
                    "page_number": None,
                    "chunk_index": i
                })

            logger.info(f"DOCX 解析完成，共切分为 {len(chunks)} 个切片")
            return chunks
        except Exception as e:
            logger.error(f"解析 DOCX 发生异常: {e}")
            raise e

    def _parse_txt(self, file_path: str) -> List[Dict[str, Any]]:
        """
        解析文本 TXT 文件。兼容 UTF-8 和 GBK 编码，没有页码，置为 None 并记录片段编号。
        """
        chunks = []
        full_text = ""
        # 尝试 UTF-8 编码读取
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                full_text = f.read()
        except UnicodeDecodeError:
            logger.warning(f"TXT 文件 UTF-8 解码失败，尝试 GBK 编码: {file_path}")
            # 尝试 GBK 编码读取
            try:
                with open(file_path, 'r', encoding='gbk') as f:
                    full_text = f.read()
            except Exception as e:
                logger.error(f"TXT 文件 GBK 编码读取也失败: {e}")
                raise e
        except Exception as e:
            logger.error(f"读取 TXT 文件失败: {e}")
            raise e

        if not full_text.strip():
            logger.warning(f"TXT 文件无内容: {file_path}")
            return []

        # 切分文本
        split_chunks = self.splitter.split_text(full_text)
        for i, c in enumerate(split_chunks):
            chunks.append({
                "text": c,
                "page_number": None,
                "chunk_index": i
            })

        logger.info(f"TXT 解析完成，共切分为 {len(chunks)} 个切片")
        return chunks
